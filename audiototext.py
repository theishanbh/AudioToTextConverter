# os tutorial
import os
import speech_recognition as sr
import audioread
from pydub import AudioSegment
import docx
from docx import Document



# f = open('.\TextVersion\hello.txt', "w")


os.chdir('AudioFiles')
filelist = os.listdir()

os.chdir('F:\Personal Project\Youtuve-dl&converter')

r = sr.Recognizer()

def duration_detector(length):
    mins = length // 60  # calculate in minutes
    length %= 60
    seconds = length  # calculate in seconds
    return mins, seconds
# 10 basic expressions with être (to be) _ Super Easy French 69-km2Pq_bSOEU
for filename in filelist:
    text_pathname = 'F:\Personal Project\Youtuve-dl&converter\TextVersion\\'+filename[:-4] + '.docx'
    file_pathname = 'F:\Personal Project\Youtuve-dl&converter\AudioFiles\\'+filename
    document = Document()
    document.save(text_pathname)
    print(filename)

    with audioread.audio_open(file_pathname) as f:
        totalsec = f.duration
        minutes, seconds = duration_detector(int(totalsec))
        print(minutes , seconds)

    for i in range(minutes):
        print(i)
        t1 = i * 60 * 1000 + 1#Works in milliseconds
        t2 = (i+1)* 60 * 1000
        newAudio = AudioSegment.from_wav(file_pathname)
        newAudio = newAudio[t1:t2]
        newAudio.export('newSong.wav', format="wav")
        tempname = 'newSong.wav'
        with sr.AudioFile(tempname) as source:
                audio_data = r.record(source)
                try:
                    text = r.recognize_google(audio_data, language='en-GB')
                except:
                    text = "Couldn't identify audio"
                document.add_paragraph(text+"\n")
        print(text)
    else:
        if seconds == 0:
            pass
        else:
            t1 = t2
            t2 = t2 + seconds*1000
            newAudio = AudioSegment.from_wav(file_pathname)
            newAudio = newAudio[t1:t2]
            newAudio.export('newSong.wav', format="wav") 
            tempname = 'newSong.wav'
            with sr.AudioFile(tempname) as source:
                audio_data = r.record(source)
                try:
                    text = r.recognize_google(audio_data, language='en-GB') # bcp 47 language tag
                except:
                    text = "Couldn't identify audio"    
                document.add_paragraph(text + "\n")
            print(text)
            
                
    document.save(text_pathname)

# cwd = os.getcwd() 
# print("Current working directory:", cwd) 

# path = "/"
# dir_list = os.listdir(path) 
  
# print("Files and directories in '", path, "' :") 
# print(dir_list) 